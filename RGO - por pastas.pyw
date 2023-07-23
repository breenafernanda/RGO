import docx
from docx.shared import Inches
import os
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from docx.shared import Cm
from PIL import Image, ImageTk
import pyautogui as pg
import tkinter as tk
from tkinter import ttk
import time
from google_drive_downloader import GoogleDriveDownloader as gdd



def janela_rgo():
    def criar_pasta_cliente(opcoes_franq, input_cliente):
        franqueado_selecionado = opcoes_franq.get()
        nome_cliente = input_cliente.get()

        caminho_pasta_franqueados = 'config/Franqueados'
        caminho_franqueado_selecionado = os.path.join(caminho_pasta_franqueados, franqueado_selecionado)

        # Verifica se a pasta do franqueado selecionado existe
        if os.path.exists(caminho_franqueado_selecionado):
            caminho_cliente = os.path.join(caminho_franqueado_selecionado, nome_cliente)
            
            # Verifica se a pasta do cliente já existe
            if not os.path.exists(caminho_cliente):
                try:
                    os.makedirs(caminho_cliente)
                    print("Pasta do cliente criada com sucesso!")

                    # Criar a pasta de fotos dentro da pasta do cliente
                    caminho_pasta_fotos = os.path.join(caminho_cliente, "Fotos")
                    os.makedirs(caminho_pasta_fotos)
                    print("Pasta de fotos criada com sucesso!")
                except OSError:
                    print("Erro ao criar pasta do cliente ou pasta de fotos.")
            else:
                print("A pasta do cliente já existe.")
        else:
            print("Pasta do franqueado não encontrada.")
    root = tk.Tk()
    root.title("Kinsol Energias Renováveis - RGO")

    # Altera o tema da interface gráfica
    root.configure(bg='#FFFFFF')  # Fundo branco

    # Cria os widgets da janela
    label0 = tk.Label(root, text="KINSOL - Energias Renováveis", bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 16, "bold"))
    tabel01 = tk.Label(root, text="Relatório de Gestão de Obras (RGO)", bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 16, "bold"))

    label1 = tk.Label(root, text="Cliente:", bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))
    label2 = tk.Label(root, text="E-mail:", bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))
    label3 = tk.Label(root, text="Endereço:", bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))
    label4 = tk.Label(root, text="Fone:", bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))

    label_telhado = tk.Label(root, text="Telhado:", bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))
    label_caibro = tk.Label(root, text="Caibro:", bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))
    label_disjuntor = tk.Label(root, text="Disjuntor:", bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))
    label_corrente = tk.Label(root, text="Corrente [A]:", bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))
    label_adequacao = tk.Label(root, text="Adequação do Padrão:", bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))
    label_franqueado = tk.Label(root, text="Franqueado:", bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))

    input_cliente = tk.Entry(root, width=45, bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))
    input_email = tk.Entry(root, width=45, bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))
    input_end = tk.Entry(root, width=45, bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))
    input_fone = tk.Entry(root, width=45, bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"))
    input_corrente = tk.Entry(root, width=7, bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12, "bold"), justify="center")

    button_rgo = tk.Button(root, text="ELABORAR RGO", command=BOTAO_RGO, bg="#A54694", fg='#FFFFFF', width=15)  # Botão roxo com texto branco e tamanho maior
    button_pasta = tk.Button(root, text="Abrir pasta Fotos", command=BOTAO_PASTA, bg="#4cb2ba")
    button_elaborar_rgo = tk.Button(root, text="Abrir pasta RGO", command=BOTAO_PASTA_RGO, bg='#4cb2ba')
    button_categorias = tk.Button(root, text="CATEGORIAS", command=BOTAO_CATEGORIAS, bg='#4cb2ba')
    button_franqueado = tk.Button(root, text="PASTA FRANQUEADOS", command=BOTAO_FRANQUEADOS, bg='#4cb2ba')
    button_apagar_fotos = tk.Button(root, text="APAGAR FOTOS", command=BOTAO_APAGAR, bg="#D8A7C1")  # Botão roxo claro
    button_criar_pasta = tk.Button(root, text="CRIAR PASTAS", command=BOTAO_CRIAR_PASTAS, bg="#D8A7C1")  # Botão roxo claro
    button_mais = tk.Button(root, text=" + ", command=BOTAO_MAIS, bg="#A54694", fg='#FFFFFF')  # Botão roxo com texto branco

    output = tk.Text(root, width=85, height=25, bg='#FFFFFF', fg='#A54694', font=("Times New Roman", 12))

    # OPÇÕES DE TELHADOS
    opcoes_telhado = tk.StringVar(root)
    opcoes_telhado.set("Cerâmico")  # Set a default option
    options = ["Cerâmico", "Fibrocimento", "Metálico", "Laje", "Solo", "Estacionamento"]
    menu_telhados = tk.OptionMenu(root, opcoes_telhado, *options, command=telhado_selecionado)
    menu_telhados.config(bg="#4cb2ba", font=("Times New Roman", 11))

    # OPÇÕES DE CAIBRO
    opcoes_caibro = tk.StringVar(root)
    opcoes_caibro.set("Madeira")  # Set a default option
    options = ["Madeira", "Metálico"]
    menu_caibros = tk.OptionMenu(root, opcoes_caibro, *options, command=caibro_selecionado)
    menu_caibros.config(bg="#4cb2ba", font=("Times New Roman", 11))

    # OPÇÕES DE DISJUNTOR
    opcoes_disjuntor = tk.StringVar(root)
    opcoes_disjuntor.set("Trifásico")  # Set a default option
    options = ["Monofásico 127V", "Monofásico 220V", "Bifásico", "Trifásico"]
    menu_disjuntor = tk.OptionMenu(root, opcoes_disjuntor, *options, command=disjuntor_selecionado)
    menu_disjuntor.config(bg="#4cb2ba", font=("Times New Roman", 11))

    # OPÇÕES DE ADEQUAÇÃO DO PADRÃO
    opcoes_adequacao = tk.StringVar(root)
    opcoes_adequacao.set("Não")  # default option
    options = ["Sim", "Não"]
    menu_adequacao = tk.OptionMenu(root, opcoes_adequacao, *options, command=adequacao_selecionado)
    menu_adequacao.config(bg="#4cb2ba", font=("Times New Roman", 11))

    # LINHA 1 - FRANQUEADO
    opcoes_franq = tk.StringVar(root)
    opcoes_franq.set("Franqueado Kinsol")  # default option

    caminho_pasta_franqueados = 'config/Franqueados'
    opt_franq = os.listdir(caminho_pasta_franqueados)

    menu_franqueado = tk.OptionMenu(root, opcoes_franq, *opt_franq, command=franqueados)
    menu_franqueado.config(bg="#4cb2ba")

    # Posiciona os widgets na janela
    label0.grid(row=0, column=0, columnspan=2)
    tabel01.grid(row=1, column=0, columnspan=2)
    label_franqueado.grid(row=0, column=2)
    menu_franqueado.grid(row=1, column=2)
    button_mais.grid(row=1, column=3)
    label1.grid(row=2, column=0)
    label2.grid(row=3, column=0)
    label3.grid(row=4, column=0)
    label4.grid(row=5, column=0)
    input_cliente.grid(row=2, column=1)
    input_email.grid(row=3, column=1)
    input_end.grid(row=4, column=1)
    input_fone.grid(row=5, column=1)
    label_telhado.grid(row=7, column=0)
    menu_telhados.grid(row=8, column=0)
    label_caibro.grid(row=7, column=1)
    menu_caibros.grid(row=8, column=1)
    label_disjuntor.grid(row=7, column=2)
    menu_disjuntor.grid(row=8, column=2)
    label_adequacao.grid(row=10, column=0)
    menu_adequacao.grid(row=11, column=0)
    label_corrente.grid(row=10, column=1, columnspan=1)
    input_corrente.grid(row=11, column=1, columnspan=1)

    output.grid(row=15, column=0, columnspan=4, padx=5, pady=5)  # Adicionando espaços antes e depois do output
    button_franqueado.grid(row=13, column=0, padx=5, pady=5)
    button_categorias.grid(row=13, column=1, padx=5, pady=5)
    button_rgo.grid(row=13, column=2, padx=5, pady=5)

    button_pasta.grid(row=17, column=0, columnspan=1, padx=5, pady=5)
    button_elaborar_rgo.grid(row=17, column=1, columnspan=1, padx=5, pady=5)
    button_apagar_fotos.grid(row=17, column=2, columnspan=1, padx=5, pady=5)
    button_criar_pasta.grid(row=17, column=3, columnspan=1, padx=5, pady=5)

    return root, input_cliente, input_email, input_end, input_fone, opcoes_telhado, opcoes_adequacao, opcoes_disjuntor, opcoes_caibro, input_corrente, opcoes_franq, output
 

def imprimir(text):
    output.insert(tk.END, f"{text}\n")
    output.see(tk.END) # auto scroll

def restart_window():
    global root  # Make 'root' a global variable so it can be accessed outside the function
    root.destroy()  # Destroy the current window


def BOTAO_MAIS():
    imprimir('criar pastas')

    def janela_adicionar_franqueado():
        root_franqueados = tk.Tk()
        root_franqueados.title("Kinsol - Adicionar Franqueado")
    
        # Altera o tema da interface gráfica
        root_franqueados.configure(bg='#A54694')

        label_branco01 = tk.Label(root_franqueados, text="  ADICIONAR FRANQUEADO",bg='#A54694',  fg='#FFFFFF',font=("Times New Roman", 18,"bold"))
        label_branco02 = tk.Label(root_franqueados, text="   ",bg='#A54694',  fg='#FFFFFF',font=("Times New Roman", 18,"bold"))
        label_franqueado = tk.Label(root_franqueados, text="Nome do Franqueado: ",bg='#A54694',  fg='#FFFFFF',font=("Times New Roman", 12,"bold"))
        label_franquead2o= tk.Label(root_franqueados, text="       ",bg='#A54694',  fg='#FFFFFF',font=("Times New Roman", 12,"bold"))
        input_franqueado = tk.Entry(root_franqueados, width=45, bg = '#FFFFFF', fg='#4cb2ba', font=("Times New Roman", 12,"bold"))
        botao_add = tk.Button(root_franqueados, text=" + ", command=BOTAO_ADICIONAR_FRANQUEADO, bg="green")

        label_branco01.grid(row=0, column=1, rowspan=3, columnspan=5)
        label_franqueado.grid(row=3, column=0)
        label_franquead2o.grid(row=3, column=2)
        input_franqueado.grid(row=3, column=1)
        botao_add.grid(row=3, column=2)
        label_branco02.grid(row=4, column=0,rowspan=3, columnspan=5)

        return root_franqueados, input_franqueado
    
    def ADICIONAR_TXT_FRANQUEADO(franqueado):
        with open('franqueados.txt', 'a',encoding='utf-8') as arquivo:
            arquivo.write(franqueado + '\n')

    def BOTAO_ADICIONAR_FRANQUEADO():
            # adicionar no bloco de notas
            franqueado = input_franqueado.get()
            ADICIONAR_TXT_FRANQUEADO(franqueado)
            
            BOTAO_CRIAR_PASTAS()
            restart_window()
            imprimir(f'Franqueado adicionado: {franqueado}')
            root_franqueados.destroy()
            os.startfile('RGO - por pastas.exe')
        
    root_franqueados, input_franqueado = janela_adicionar_franqueado()

    root_franqueados.geometry("+1000+600")
    
    root_franqueados.mainloop()  # Start the main event loop of the window[]

def BOTAO_CRIAR_PASTAS():
    # Caminho do arquivo franqueados.txt
    caminho_arquivo = 'franqueados.txt'
    
    # Verifica se o arquivo existe
    if os.path.isfile(caminho_arquivo):
        # Lê o arquivo franqueados.txt
        with open(caminho_arquivo, 'r', encoding='utf-8') as arquivo:
            linhas = arquivo.readlines()
        # Remove espaços em branco e quebras de linha das linhas
        linhas = [linha.strip() for linha in linhas]
        # print(linhas)
        

        # Cria as pastas com base nas linhas do arquivo
        for linha in linhas:
            nome_pasta = linha.replace('/', '_')  # Substitui '/' por _
            nome_pasta = linha.replace(' ', '_')  # Substitui ' ' por _
            
            try:
                criar_pasta = f'config\Franqueados\{nome_pasta}'
                os.mkdir(criar_pasta)
                imprimir(f"Pasta criada: {nome_pasta}")
            except FileExistsError:
                print(f"Pasta já existe: {nome_pasta}")
    else:
        print("O arquivo franqueados.txt não existe.")
    root.update()

def franqueados():
    imprimir(f"Franqueado Selecionado: {opcoes_telhado}")

def BOTAO_FRANQUEADOS():
    #abrir a pasta que está no mesmo diretório e chama Fotos
    folder_path = os.path.join("config\Franqueados")
    os.system("start " + folder_path)

def BOTAO_CATEGORIAS():
    output.delete("1.0", tk.END)
    checklist = f"CATEGORIAS DE FOTOS DO RGO:\n\n1. Frontal do Inversor\n2. Acabamentos CA\n3. Stick do inversor\n4. Etiqueta da lateral do inversor\n5. Ponto de conexão com a rede\n6. Quadro de disjuntores\n7. Disjuntor do inversor\n8. Transformador (se necessário)\n9. Padrão do Cliente\n10. Aterramento dos módulos e da carcaça do inversor\n11. Placa de Geração Própria\n12. Placa da Kinsol\n13. Medições CA\n14. Medições CC\n15. Desenho CROQUI\n16. Telhado ANTES da instalação\n17. Telhado DEPOIS da instalação\n18. Vedações do Telhado\n19. Fachada do cliente\n20. Kit do Sistema"
    imprimir(checklist)

def BOTAO_APAGAR():
    imprimir('apagar fotos')

  # LISTAR FOTOS DA PASTA E ADICIONAR NO RELATORIO
    caminho_base = 'config\pictures' 
    lista_de_arquivos = os.listdir(caminho_base)
    # imprimir(lista_de_arquivos)
    # imprimir(categorias)
    for pasta in lista_de_arquivos:
        # print(pasta) # PASTA TEM O NOME DA CATEGORIA

        caminho_pasta = f'config\pictures\{pasta}' 

        # listar fotos dentro da pasta
        lista_de_fotos = os.listdir(caminho_pasta)
        # imprimir (lista_de_fotos)
        contador = 0
        for foto in lista_de_fotos:
            caminho_da_foto = f'{caminho_pasta}\{foto}'
            os.remove(caminho_da_foto)
            # imprimir(f'apagar foto -> {pasta} -> {foto}')
    output.delete("1.0", tk.END)
    imprimir('TODAS AS FOTOS FORAM APAGADAS...')
    # time.sleep(3)
    # BOTAO_CATEGORIAS()



def BOTAO_PASTA():
     #abrir a pasta que está no mesmo diretório e chama Fotos
    folder_path = os.path.join("config\pictures")
    os.system("start " + folder_path)

def BOTAO_PASTA_RGO():
    # Obtenha o diretório atual (onde o arquivo .py está localizado)
    current_directory = os.getcwd()
    # Crie o caminho completo para a pasta
    folder_path = os.path.join("config\RGOs")
    os.system("start " + folder_path)

def BOTAO_PASTA_cliente(franqueado,cliente):
    # Crie o caminho completo para a pasta

    folder_path = os.path.join(f'config\Franqueados\{franqueado}\{cliente}')
    # folder_path=f'"{folder_path}"'
    imprimir(folder_path)
    try:
        os.system("start " + folder_path)
    except Exception as e: imprimir(e)

# Criar função para imprimir no output o texto
def imprimir(text):
    try:
        output.insert(tk.END, f"{text}\n")
        output.see(tk.END) # auto scroll
    except Exception as e: print('')

def telhado_selecionado(opcoes_telhado):

    # output.delete("1.0", tk.END)
    imprimir(f"Tipo de Telhado Selecionado: {opcoes_telhado}")

def caibro_selecionado(opcoes_caibro):
    # output.delete("1.0", tk.END)
    imprimir(f"Tipo de Caibro Selecionado: {opcoes_caibro}")

def disjuntor_selecionado(opcoes_disjuntor):
    # output.delete("1.0", tk.END)
    imprimir(f"Tipo de Disjuntor Selecionado: {opcoes_disjuntor}")

def adequacao_selecionado(opcoes_adequacao):
    # output.delete("1.0", tk.END)
    imprimir(f"Tipo de Adequação de Padrão Selecionado: {opcoes_adequacao}")



def BOTAO_RGO():
    #verificar se o relatorio_RGO.docx existe
        #se não existir, realizar o download com o modulo googledrivedownloader e unzip no diretorio correto
    # Abrir modelo de RGO

    # Caminho para o diretório onde o arquivo será salvo
    diretorio = 'config'

    # Nome do arquivo a ser verificado e baixado
    nome_arquivo = 'relatorio_RGO.docx'

    # Verificar se o arquivo existe
    caminho_arquivo = os.path.join(diretorio, nome_arquivo)
    if not os.path.exists(caminho_arquivo):
        # Realizar o download usando o GoogleDriveDownloader
        file_id = 'COLOQUE_AQUI_O_ID_DO_ARQUIVO_NO_GOOGLE_DRIVE'
        gdd.download_file_from_google_drive(file_id=file_id, dest_path=caminho_arquivo, overwrite=True)

    # Abrir o modelo de RGO com a biblioteca docx
    # document = docx.Document(caminho_arquivo)
    document = docx.Document(r'config\relatorio_RGO.docx')
    

    def salvar(cliente):
        # imprimir('\n')
        
        try:
            document.save(f'config\RGOs\Relatório de Gestão de Obra_RGO_{cliente}.docx')


        except: 
            # imprimir(f'Erro ao salvar, fechar arquivo e tentar novamente')
            try:
                i = 1
                aberto = True
                while (aberto==True):
                    try:
                        # document.save(f'config\RGOs\Relatório de Gestão de Obra_RGO_{cliente}({i}).docx')
                        imprimir(f'Documento salvo como:')
                        # imprimir(f'Relatório de Gestão de Obra_RGO_{cliente}({i}).docx')
                        aberto = '✘'
        #                 i=0
                    except:
                        i = i + 1
            except Exception as e: print(e)

    def salvar_backup(cliente):
        # imprimir('\n')
        
        try:
            document.save(f'config\Franqueados\{franqueado}\{cliente}\Relatório de Gestão de Obra_RGO_{cliente}.docx')
            imprimir(f'Documento salvo na pasta {cliente} como:')
            imprimir(f'Relatório de Gestão de Obra_RGO_{cliente}.docx')

        except: 
            # imprimir(f'Erro ao salvar, fechar arquivo e tentar novamente')
            try:
                i = 1
                aberto = True
                while (aberto==True):
                    try:
                        document.save(f'config\Franqueados\{franqueado}\{cliente}\Relatório de Gestão de Obra_RGO_{cliente}({i}).docx')
                        imprimir(f'Documento salvo na pasta {cliente} como:')
                        imprimir(f'Relatório de Gestão de Obra_RGO_{cliente}({i}).docx')
                        aberto = '✘'
        #                 i=0
                    except:
                        i = i + 1
            except Exception as e: print(e)

    def obter_tamanho(imagem):
        try:
            # obter as dimensões da imagem
            img = Image.open(imagem)
            largura, altura = img.size
            if largura > altura: tamanho = 3
            else: tamanho =2
            # print(tamanho)
        except:
            tamanho = 3
        return tamanho

    def obter_foto(tipo_de_foto):
        # print(tipo_de_foto)
        if tipo_de_foto == 1:
            foto = 'Frontal do Inversor'
        if tipo_de_foto == 2:
            foto = 'Acabamentos CA'
        if tipo_de_foto == 3:
            foto = 'Etiqueta do Stick'
        if tipo_de_foto == 4:
            foto = 'Etiqueta da Lateral do Inversor'
        if tipo_de_foto == 5:
            foto = 'Ponto de Conexão'
        if tipo_de_foto == 6:
            foto = 'Caixa de Disjuntores'
        if tipo_de_foto == 7:
            foto = 'Disjuntor do Inversor'
        if tipo_de_foto == 8:
            foto = 'Transformador'
        if tipo_de_foto == 9:
            foto = 'Padrão do Cliente'
        if tipo_de_foto == 10:
            foto = 'Aterramento da Carcaça e dos módulos'
        if tipo_de_foto == 11:
            foto = 'Placa de Advertência'
        if tipo_de_foto == 12:
            foto = 'Placa da Kinsol'
        if tipo_de_foto == 13:
            foto = 'Medições CA'
        if tipo_de_foto == 14:
            foto = 'Medições CC'
        if tipo_de_foto == 15:
            foto = 'Desenho Croqui'
        if tipo_de_foto == 16:
            foto = 'Telhado Antes'
        if tipo_de_foto == 17:
            foto = 'Telhado Depois'
        if tipo_de_foto == 18:
            foto = 'Vedações'
        if tipo_de_foto == 19:
            foto = 'Fachada do Cliente'
        if tipo_de_foto == 20:
            foto = 'Kit do sistema'
        return str(foto)

    def ordenar(lista):
        # print(lista)

        #Ordenar lista numericamente
        try:
            lista_ordenada = sorted(lista, key=lambda x: int(x.split()[1][:-4]))
            return lista_ordenada

        except: 
            return lista

    def ordenar_categorias(lista):
        # lista = ['1.jpg', '10 1.jpg', '10 2.jpg', '11.jpg', '12.jpg', '13 1.jpg', '13 2.jpg', '13 3.jpg', '14 1.jpg', '14 10.jpg', '14 11.jpg', '14 12.jpg', '14 13.jpg', '14 14.jpg', '14 15.jpg', '14 2.jpg', '14 3.jpg', '14 4.jpg', '14 5.jpg', '14 6.jpg', '14 7.jpg', '14 8.jpg', '14 9.jpg', '15.jpg', '16.jpg', '17.jpeg', '18.jpg', '19.jpg', '2.jpg', '3 1.jpg', '3 2.jpg', '4 1.jpg', '4 2.jpg', '5.jpg', '6.jpg', '7.jpg', '8.jpg', '9.jpg']

        # Extrair o número de cada arquivo e ordenar a lista
        lista_ordenada = sorted(lista, key=lambda x: int(x.split()[0]) if ' ' in x else int(x.split('.')[0]))

        return lista_ordenada

    def separar_fotos_por_categoria(lista_de_arquivos):
        categorias = {}
        
        for arquivo in lista_de_arquivos:
            categoria = re.search(r'^(\d+)', arquivo).group(1)
            categoria_descricao = obter_foto(int(categoria))
            
            if categoria_descricao not in categorias:
                categorias[categoria_descricao] = []
            
            categorias[categoria_descricao].append(arquivo)
            
        return categorias

    def inserir_foto(cell, foto):
        tamanho = obter_tamanho(foto)
        imagem = cell.add_paragraph().add_run()
        imagem.add_picture(foto, width=Inches(tamanho))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def obter_cell(categoria, contador):
        if categoria == 'Frontal do Inversor':
                table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de aparição
                celula = table.cell(3, 2) # índices começando em 0
        if categoria == 'Acabamentos CA':
                table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de aparição
                celula = table.cell(3, 8) # índices começando em 0
        if categoria == 'Etiqueta do Stick':       
                table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de aparição
                celula = table.cell(5, contador-1) # índices começando em 0
        if categoria == 'Etiqueta da Lateral do Inversor':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(5, 8) # índices começando em 0
        if categoria == 'Ponto de Conexão':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(8, 0) # índices começando em 0
        if categoria == 'Caixa de Disjuntores':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(10, 0) # índices começando em 0
        if categoria == 'Disjuntor do Inversor':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(10, 8) # índices começando em 0
        if categoria == 'Transformador':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(12, 0) # índices começando em 0
        if categoria =='Padrão do Cliente':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(14, 0) # índices começando em 0
        if categoria == 'Aterramento da Carcaça e dos módulos':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(16, contador) # índices começando em 0
        if categoria == 'Placa de Advertência':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(19, 0) # índices começando em 0
        if categoria =='Placa da Kinsol':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(19, 8) # índices começando em 0
        if categoria ==  'Medições CA':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            if contador == 1: contador = 0
            if contador == 2: contador = 5
            if contador == 3: contador = 8
            #    print(contador, 'ca')
            celula = table.cell(23, contador) # índices começando em 0
        if categoria ==  'Medições CC':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            
            if contador > 3:

                    if contador == 4:
                        contador = 0
                        celula = table.cell(28, contador) # índices começando em 0

                    if contador == 5:
                        contador = 5
                        celula = table.cell(28, contador) # índices começando em 0

                    if contador == 6:
                        contador == 8
                        celula = table.cell(28, contador) # índices começando em 0

                    if contador == 7:
                        contador = 0
                        celula = table.cell(30, contador) # índices começando em 0

                    if contador ==8:
                        contador = 5
                        celula = table.cell(30, contador) # índices começando em 0

                    if contador == 9:
                        contador = 8
                        # print('contador 9', contador)
                        celula = table.cell(30, contador) # índices começando em 0

                    if contador == 10:
                        contador = 1
                        celula = table.cell(32, contador) # índices começando em 0

                    if contador == 11:
                        contador = 5
                        celula = table.cell(32, contador) # índices começando em 0

                    if contador == 12:
                        contador = 8
                        celula = table.cell(32, contador) # índices começando em 0

                    if contador == 13:
                        contador = 0
                        celula = table.cell(34, contador) # índices começando em 0

                    if contador ==14:
                        contador = 5
                        celula = table.cell(34, contador) # índices começando em 0

                    if contador == 15:
                        contador = 8
                        celula = table.cell(34, contador) # índices começando em 0

            else:
                    # print( contador, 'cc')
                    if contador == 2: contador = 5
                    if contador == 3:
                        contador = 8
                    celula = table.cell(26, contador) # índices começando em 0
            
            #    print(contador, 'CC')
        if categoria == 'Desenho Croqui':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(37, 7) # índices começando em 0
        if categoria == 'Telhado Antes':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(39, 0) # índices começando em 0
        if categoria ==  'Telhado Depois':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(41, 0) # índices começando em 0
        if categoria == 'Vedações':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(43, 0) # índices começando em 0
        if categoria == 'Fachada do Cliente':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de apariçã
            celula = table.cell(45, 0) # índices começando em 0
        if categoria == 'Kit do sistema':
            table = document.tables[0] # assume que a tabela desejada é a primeira na ordem de aparição
            celula = table.cell(10, 0) # índices começando em 0

        return table, celula

    def adicionar_foto(categoria,contador, arquivo, caminho_pasta):
        # imprimir(f'{categoria}, foto:{contador}, arquivo: {arquivo}')
        foto = f'{caminho_pasta}/{arquivo}'
        # print(foto)
        table, celula = obter_cell(categoria, contador)
        
        inserir_foto(celula, foto)

    def adicionar_dados_cadastro(dado):
        # obter célula de Nome
        table = document.tables[0] # assume que a tabela desejada é a primeira na ordem de aparição
        if dado == 'nome':
            cell = table.cell(1,2) # índices começando em 0
        if dado == 'email':
            cell = table.cell(2,2)
        if dado == 'end':
            cell = table.cell(3,5)
        if dado == 'fone':
            cell = table.cell(2,15)

        return table, cell

    def inserir_informacoes(valor_telhado,valor_adequacao,valor_disjuntor, valor_caibro, corrente):
        try: # TELHADO
            if valor_telhado == 'Cerâmico':
                # imprimir(valor_telhado)
                
                table = document.tables[0] # definir a tabela
                cell = table.cell(5, 4) 
                cell.text = ' ☑    Cerâmico'
            if valor_telhado == 'Fibrocimento':
                # imprimir(valor_telhado)
                
                table = document.tables[0] # definir a tabela
                cell = table.cell(5, 6) 
                cell.text = ' ☑   Fibrocimento'
            if valor_telhado == 'Laje':
                # imprimir(valor_telhado)
                
                table = document.tables[0] # definir a tabela
                cell = table.cell(5, 15) 
                cell.text = ' ☑  Laje'
            if valor_telhado == 'Metálico':
                # imprimir(valor_telhado)
                
                table = document.tables[0] # definir a tabela
                cell = table.cell(5, 10) 
                cell.text = ' ☑  Metálico'
            if valor_telhado == 'Solo':
                # imprimir(valor_telhado)
                
                table = document.tables[0] # definir a tabela
                cell = table.cell(5, 20) 
                cell.text = ' ☑  Solo'
            if valor_telhado == 'Estacionamento':
                # imprimir(valor_telhado)
                
                table = document.tables[0] # definir a tabela
                cell = table.cell(5, 25) 
                cell.text = ' ☑    Estacionamento'
        except Exception as E: print(E)
        try: # ADEQUAÇÃO DO PADRÃO
            if valor_adequacao == 'Sim':
                table = document.tables[0] # definir a tabela
                cell = table.cell(6, 6) 
                cell.text = ' ☑    Sim'

            if valor_adequacao == 'Não':
                table = document.tables[0] # definir a tabela
                cell = table.cell(6, 10) 
                cell.text = ' ☑    Não'           
        except Exception as E: print(E)
        try: # CAIBRO
            if valor_caibro == 'Madeira':
                table = document.tables[0] # definir a tabela
                cell = table.cell(6, 16) 
                cell.text = ' ☑   Madeira'
            if valor_caibro == 'Metálico':
                table = document.tables[0] # definir a tabela
                cell = table.cell(6, 20) 
                cell.text = ' ☑   Metálico'
        except Exception as e: print(e)
        try: # DISJUNTOR DO PADRÃO
                if valor_disjuntor == 'Monofásico 127V':
                    table = document.tables[0] # definir a tabela
                    cell = table.cell(7, 5) 
                    cell.text = ' ☑  Monofásico         127'
                if valor_disjuntor == 'Monofásico 220V':
                    table = document.tables[0] # definir a tabela
                    cell = table.cell(7, 5) 
                    cell.text = ' ☑  Monofásico         220'
                if valor_disjuntor == 'Bifásico':
                    table = document.tables[0] # definir a tabela
                    cell = table.cell(7, 10) 
                    cell.text = ' ☑   Bifásico'  
                if valor_disjuntor == 'Trifásico':
                    table = document.tables[0] # definir a tabela
                    cell = table.cell(7, 15) 
                    cell.text = ' ☑   Trifásico'                      
        except Exception as e:print(e)  
        try: # CORRENTE
                    table = document.tables[0] # definir a tabela
                    cell = table.cell(7, 23) 
                    cell.text = corrente
        except Exception as e: print(e)
            
    def inserir_dados(dado, texto):
        # inserir texto na cell da table 
        
        # inserir nome 
        table, cell = adicionar_dados_cadastro(dado)
        try:
            cell.text = texto
        except: 
            imprimir(f'ERRO ao inserir {dado} - > {texto}')


        

        # Endereço
        # Telefone

    def preencher_checklist(categoria):
            # imprimir(categoria)

            if categoria == 'Frontal do Inversor':
                table = document.tables[0] # definir a tabela
                # Alinhe o conteúdo centralizado na célula
                cell = table.cell(19, 0) 
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.text = '✔'

            if categoria == 'Acabamentos CA':
                table = document.tables[0] # definir a tabela
                cell = table.cell(20, 0) 
                cell.text = '✔'
                acabamento = '✔'
            if categoria == 'Etiqueta do Stick':
                table = document.tables[0] # definir a tabela
                cell = table.cell(21, 0) 
                cell.text = '✔'
            if categoria == 'Etiqueta da Lateral do Inversor':
                table = document.tables[0] # definir a tabela
                cell = table.cell(22, 0) 
                cell.text = '✔'
            if categoria == 'Ponto de Conexão':
                table = document.tables[0] # definir a tabela
                cell = table.cell(23, 0) 
                cell.text = '✔'
            if categoria == 'Caixa de Disjuntores':
                table = document.tables[0] # definir a tabela
                cell = table.cell(24, 0) 
                cell.text = '✔'
            if categoria == 'Disjuntor do Inversor':
                table = document.tables[0] # definir a tabela
                cell = table.cell(25, 0) 
                cell.text = '✔'
            if categoria == 'Transformador':
                table = document.tables[0] # definir a tabela
                cell = table.cell(26, 0) 
                cell.text = '✔'
            if categoria == 'Padrão do Cliente':
                table = document.tables[0] # definir a tabela
                cell = table.cell(27, 0) 
                cell.text = '✔'
            if categoria == 'Aterramento da Carcaça e dos módulos':
                table = document.tables[0] # definir a tabela
                cell = table.cell(28, 0) 
                cell.text = '✔'
                
            if categoria == 'Placa de Advertência':
                table = document.tables[0] # definir a tabela
                cell = table.cell(19, 10) 
                cell.text = '✔'
            if categoria == 'Placa da Kinsol':
                table = document.tables[0] # definir a tabela
                cell = table.cell(20, 10) 
                cell.text = '✔'
            if categoria == 'Medições CA':
                table = document.tables[0] # definir a tabela
                cell = table.cell(21, 10) 
                cell.text = '✔'
            if categoria == 'Medições CC':
                table = document.tables[0] # definir a tabela
                cell = table.cell(22, 10) 
                cell.text = '✔'
            if categoria == 'Desenho Croqui':
                table = document.tables[0] # definir a tabela
                cell = table.cell(23, 10) 
                cell.text = '✔'
            if categoria == 'Telhado Antes':
                table = document.tables[0] # definir a tabela
                cell = table.cell(24, 10) 
                cell.text = '✔'
            if categoria == 'Telhado Depois':
                table = document.tables[0] # definir a tabela
                cell = table.cell(25, 10) 
                cell.text = '✔'
            if categoria == 'Vedações':
                table = document.tables[0] # definir a tabela
                cell = table.cell(26, 10) 
                cell.text = '✔'
            if categoria == 'Fachada do Cliente':
                table = document.tables[0] # definir a tabela
                cell = table.cell(27, 10) 
                cell.text = '✔'

            return 
    
    def renomear(caminho_pasta, foto, new_name):
        try:
            caminho_antigo = f'{caminho_pasta}\{foto}'
            caminho_novo = f'{caminho_pasta}\{new_name}'
            # print(caminho_novo)
            os.rename(caminho_antigo, caminho_novo)
            novo_nome = new_name
        except Exception as e: novo_nome = foto
        return novo_nome

    def salvar_foto_backup(pasta_origem, pasta_destino, nome_arquivo):
            import shutil
            import os
            caminho_origem = os.path.join(pasta_origem, nome_arquivo)
            caminho_destino = os.path.join(pasta_destino, nome_arquivo)
            print(caminho_origem, caminho_destino)
            shutil.copy(caminho_origem, caminho_destino)
    
    output.delete("1.0", tk.END)

    # EXECUTAR PROCEDIMENTOS
    try:
        imprimir(f'Relatório de Gestão de Obras -  RGO\n')


        # DADOS DO CLIENTE  NOME, EMAIL, ENDEREÇO E TELEFONE
        try:
            try:
                cliente = input_cliente.get()
                # adicionar o nome do cliente dentro do RGO
                inserir_dados('nome', cliente)
                if cliente == "":
                    cliente = 'Cliente Kinsol'
            except: cliente = 'Cliente Kinsol'
            
            try:
                # substituir os espaços " " por "_"
               
                cliente_formatado = cliente.replace(" ", "_")
                print(cliente_formatado)
                cliente = str(cliente_formatado)
            except Exception as e:
                print(e)

            try:
                franqueado = opcoes_franq.get()
                # criar pasta do cliente, dentro da pasta do franqueado
                franqueado = franqueado.replace(' ', '_')  # Substitui ' ' por _
                caminho_pasta_franqueados = f'config\Franqueados\{franqueado}'
                pasta_cliente = f'{caminho_pasta_franqueados}\{cliente}'
                os.mkdir(pasta_cliente)
            except Exception as e: print('erro ao criar pasta cliente', e)
            try:
                caminho_pasta_fotos = f'config\Franqueados\{franqueado}\{cliente}\Fotos'
                print('aqui')
                os.mkdir(caminho_pasta_fotos)
                print('aqui2')

            except Exception as e: print(e)

            try:
                email = input_email.get()
                inserir_dados('email', email)
            except Exception as e: print(e)

            try:
                endereco = input_end.get()
                # adicionar o nome do cliente dentro do RGO
                inserir_dados('end', endereco)
            except Exception as e: print(e)
            
            try:
                telefone = input_fone.get()
                inserir_dados('fone', telefone)
            except Exception as e: print(e)
        except Exception as e: print(e)

        # INFORMAÇÕES GERAIS PARA INSTALAÇÃO DO SISTEMA
        try:
            valor_telhado = opcoes_telhado.get()
            valor_adequacao = opcoes_adequacao.get()
            valor_caibro = opcoes_caibro.get()
            valor_disjuntor = opcoes_disjuntor.get()
            corrente = input_corrente.get()
            # inserir tipo de telhado
            inserir_informacoes(valor_telhado,valor_adequacao,valor_disjuntor, valor_caibro, corrente)
        except Exception as E: print(E)

        # LISTAR FOTOS DA PASTA E ADICIONAR NO RELATORIO
        caminho_base = 'config\pictures' 
        lista_de_arquivos = os.listdir(caminho_base)
        lista_de_arquivos = ordenar_categorias(lista_de_arquivos)
        # imprimir(lista_de_arquivos)

        categorias = separar_fotos_por_categoria(lista_de_arquivos)
        # imprimir(categorias)
        try:
            for categoria, pasta in categorias.items():
                # imprimir(pasta) # PASTA TEM O NOME DA CATEGORIA
                # pasta = # somente o primeiro item da lista pasta
                for item in pasta:
                    caminho_pasta = f'config\pictures\{item}' 

                # listar fotos dentro da pasta
                lista_de_fotos = os.listdir(caminho_pasta)
                # imprimir (lista_de_fotos)
                contador = 0
                for foto in lista_de_fotos:
                    # salvar_backup(cliente,franqueado, foto)
                    contador = contador + 1
                    # new_name = renomeia_foto(categoria, contador, foto)
                    # print(new_name)
                    
                    new_name = f'{categoria} - {contador}.png'
                    try:
                        foto_nome = renomear(caminho_pasta, foto, new_name)
                        # print(new_name)
                        adicionar_foto(categoria, contador, foto_nome, caminho_pasta)
                        preencher_checklist(categoria)
                    except Exception as e: print(e)
                    try:
                        salvar_foto_backup(caminho_pasta, caminho_pasta_fotos, foto_nome)
                    except Exception as erro: print(erro)
                # imprimir(f'{categoria} -> {contador} fotos')
                if contador >= 1:
                    imprimir (f'   ✔  {categoria}')
                else:
                    imprimir (f'   Ｘ  {categoria}')
        except: imprimir (f'   Ｘ  {categoria}')
        # print(categorias)
        # for categoria, fotos in categorias.items():
        #     preencher_checklist(categoria)
        #     contador = 0
        #     fotos = ordenar(fotos)
        #     for arquivo in fotos:
        #         contador = contador + 1
        #         adicionar_foto(categoria, contador, arquivo)
        #     imprimir(f'    {contador} fotos  | {categoria} ')

        salvar(f'{cliente}')
        salvar_backup(cliente)

    except Exception as e: imprimir(e)

    BOTAO_PASTA_cliente(franqueado, cliente)








checklist = f"ANTES DE ELABORAR O RGO, RENOMEIE AS FOTOS CONFORME AS CATEGORIAS ABAIXO\n1. Frontal do Inversor\n2. Acabamentos CA\n3. Stick do inversor\n4. Etiqueta da lateral do inversor\n5. Ponto de conexão com a rede\n6. Quadro de disjuntores\n7. Disjuntor do inversor\n8. Transformador (se necessário)\n9. Padrão do Cliente\n10. Aterramento dos módulos e da carcaça do inversor\n11. Placa de Geração Própria\n12. Placa da Kinsol\n13. Medições CA\n14. Medições CC\n15. Desenho CROQUI\n16. Telhado ANTES da instalação\n17. Telhado DEPOIS da instalação\n18. Vedações do Telhado\n19. Fachada do cliente\n20. Kit do Sistema"

# Define a janela e os widgets
root, input_cliente, input_email, input_end, input_fone,opcoes_telhado, opcoes_adequacao, opcoes_disjuntor, opcoes_caibro, input_corrente ,opcoes_franq , output = janela_rgo()

root.geometry("+930+00")

imprimir(checklist)

# Inicia o loop principal da interface gráfica
root.mainloop()
